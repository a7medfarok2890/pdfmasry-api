"""Repair Arabic text direction after a pdf2docx conversion.

pdf2docx is deliberately kept as the layout engine because it preserves
tables and images well.  Some PDFs produced by browsers, however, store RTL
glyphs in logical order while positioning Arabic-Indic numbers right-to-left.
The resulting DOCX then contains reversed dates and amounts.  This module
uses the original glyph coordinates to restore those numeric tokens, applies
the OOXML RTL flags that Word expects, and recovers invoice total lines that
pdf2docx may discard as overlapping text.
"""
from __future__ import annotations

import os
import re
import unicodedata
import uuid
from collections import Counter
from collections.abc import Iterator

import fitz
import pdfplumber
from bidi.algorithm import get_display
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.table import _Cell, Table
from docx.text.paragraph import Paragraph


ARABIC_RE = re.compile(r"[\u0600-\u06ff]")
ARABIC_LETTER = r"[\u0621-\u064a]"
ARABIC_MARK = r"[\u064b-\u065f\u0670]"
ARABIC_NUMBER_RE = re.compile(
    r"(?<![\u0660-\u0669])"
    r"[\u0660-\u0669](?:[\u0660-\u0669\u066b\u066c/.,:%+-]*[\u0660-\u0669])?"
    r"(?![\u0660-\u0669])"
)
TOTAL_WORDS = ("الإجمالي", "الاجمالي", "الإجمالى", "المجموع", "المستحق")


def _dedupe_and_order_marks(match: re.Match[str]) -> str:
    """Return one canonically ordered copy of every combining mark."""
    marks = sorted(set(match.group(0)), key=unicodedata.combining)
    return "".join(marks)


def repair_arabic_combining_marks(text: str) -> str:
    """Repair common zero-width glyph artefacts emitted by PyMuPDF.

    Browser-generated PDFs can expose a combining mark immediately before
    its base letter.  A duplicated zero-width glyph may also appear beside
    the visible glyph.  The transformations below are intentionally narrow:
    normal Arabic words with correctly placed marks are left unchanged.
    """
    # A mark at the start of a word belongs to the following base letter.
    text = re.sub(
        rf"(^|[^\u0600-\u06ff])({ARABIC_MARK}+)({ARABIC_LETTER})",
        lambda match: match.group(1) + match.group(3) + match.group(2),
        text,
    )

    # In browser PDFs, marks placed on the right edge of the next glyph may
    # be reported after a hamza carrier.  Move them onto that following glyph.
    text = re.sub(
        rf"([ؤئءأإ])({ARABIC_MARK}+)({ARABIC_LETTER})",
        lambda match: match.group(1) + match.group(3) + match.group(2),
        text,
    )

    text = re.sub(rf"{ARABIC_MARK}+", _dedupe_and_order_marks, text)

    # Collapse the precise pattern produced when a visible marked glyph is
    # accompanied by a duplicate zero-width copy of the same glyph.
    text = re.sub(
        rf"({ARABIC_LETTER})({ARABIC_MARK}+)\1\1",
        lambda match: match.group(1) + match.group(2) + match.group(1),
        text,
    )
    return unicodedata.normalize("NFC", text)


def extract_numeric_replacements(pdf_path: str) -> dict[str, str]:
    """Map stored Arabic-number tokens to their visual left-to-right value.

    This is safer than reversing every number: PDFs that already store a
    token correctly simply map it to itself.
    """
    replacements: dict[str, str] = {}
    with fitz.open(pdf_path) as document:
        for page in document:
            for block in page.get_text("rawdict").get("blocks", []):
                for line in block.get("lines", []):
                    chars = [
                        char
                        for span in line.get("spans", [])
                        for char in span.get("chars", [])
                    ]
                    stored = "".join(char.get("c", "") for char in chars)
                    for match in ARABIC_NUMBER_RE.finditer(stored):
                        token_chars = chars[match.start() : match.end()]
                        visual = "".join(
                            char.get("c", "")
                            for char in sorted(
                                token_chars,
                                key=lambda char: (
                                    float(char.get("bbox", (0, 0, 0, 0))[0]),
                                    float(char.get("bbox", (0, 0, 0, 0))[1]),
                                ),
                            )
                        )
                        replacements[match.group(0)] = visual
    return replacements


def repair_extracted_text(text: str, numeric_replacements: dict[str, str]) -> str:
    """Repair a DOCX text run without disturbing Latin text or styling."""
    text = ARABIC_NUMBER_RE.sub(
        lambda match: numeric_replacements.get(match.group(0), match.group(0)),
        text,
    )
    return repair_arabic_combining_marks(text)


def extract_logical_lines(pdf_path: str) -> list[str]:
    """Return source lines in logical reading order."""
    lines: list[str] = []
    with pdfplumber.open(pdf_path) as document:
        for page in document.pages:
            for visual_line in (page.extract_text() or "").splitlines():
                logical = repair_arabic_combining_marks(
                    get_display(visual_line, base_dir="L")
                ).strip()
                if logical:
                    lines.append(logical)
    return lines


def extract_missing_total_lines(pdf_path: str) -> list[str]:
    """Extract logical Arabic invoice totals from the source PDF.

    pdfplumber exposes the physical left-to-right line; applying the Unicode
    bidirectional algorithm reconstructs its logical reading order.  Only
    numeric total/amount-summary lines are considered, which avoids adding
    ordinary page text that pdf2docx intentionally grouped elsewhere.
    """
    totals: list[str] = []
    for logical in extract_logical_lines(pdf_path):
        if (
            ARABIC_NUMBER_RE.search(logical)
            and any(word in logical for word in TOTAL_WORDS)
            and logical not in totals
        ):
            totals.append(logical)
    return totals


def _iter_cell_paragraphs(cell: _Cell) -> Iterator[Paragraph]:
    for paragraph in cell.paragraphs:
        yield paragraph
    for table in cell.tables:
        yield from _iter_table_paragraphs(table)


def _iter_table_paragraphs(table: Table) -> Iterator[Paragraph]:
    seen_cells: set[object] = set()
    for row in table.rows:
        for cell in row.cells:
            # Keep the XML element itself.  ``cell._tc`` proxies are short
            # lived, so storing only ``id(cell._tc)`` lets Python reuse an id
            # and can incorrectly skip a different cell later in the table.
            cell_element = cell._tc
            if cell_element in seen_cells:
                continue
            seen_cells.add(cell_element)
            yield from _iter_cell_paragraphs(cell)


def _iter_all_paragraphs(document) -> Iterator[Paragraph]:
    yield from document.paragraphs
    for table in document.tables:
        yield from _iter_table_paragraphs(table)
    for section in document.sections:
        yield from section.header.paragraphs
        yield from section.footer.paragraphs


def _ensure_child(parent, tag: str):
    child = parent.find(qn(tag))
    if child is None:
        child = OxmlElement(tag)
        parent.append(child)
    return child


def _mark_paragraph_rtl(paragraph: Paragraph) -> None:
    if not ARABIC_RE.search(paragraph.text):
        return

    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    paragraph_properties = paragraph._p.get_or_add_pPr()
    _ensure_child(paragraph_properties, "w:bidi")

    for run in paragraph.runs:
        run_properties = run._r.get_or_add_rPr()
        _ensure_child(run_properties, "w:rtl")
        language = _ensure_child(run_properties, "w:lang")
        language.set(qn("w:bidi"), "ar-SA")


def _mark_table_rtl(table: Table) -> None:
    if any(ARABIC_RE.search(cell.text) for row in table.rows for cell in row.cells):
        _ensure_child(table._tbl.tblPr, "w:bidiVisual")
    for row in table.rows:
        for cell in row.cells:
            for nested in cell.tables:
                _mark_table_rtl(nested)


def _normalized_for_compare(text: str) -> str:
    return re.sub(r"[\W_]+", "", text, flags=re.UNICODE)


def _character_signature(text: str) -> Counter[str]:
    return Counter(_normalized_for_compare(text))


def _align_paragraph_to_source_lines(
    paragraph: Paragraph,
    logical_lines: list[str],
) -> bool:
    """Restore line boundaries and logical order without flattening styles.

    pdf2docx sometimes concatenates adjacent source lines and may keep two
    RTL spans in physical rather than logical order.  We match by character
    multiset, then place each source line in the first run of its matching
    group.  Requiring at least two matched lines keeps this conservative and
    prevents ordinary one-line paragraphs from being rewritten.
    """
    paragraph_signature = _character_signature(paragraph.text)
    if not paragraph_signature:
        return False

    matched_lines: list[str] | None = None
    for start in range(len(logical_lines)):
        combined = Counter()
        for end in range(start, min(len(logical_lines), start + 8)):
            combined.update(_character_signature(logical_lines[end]))
            if end > start and combined == paragraph_signature:
                matched_lines = logical_lines[start : end + 1]
                break
        if matched_lines is not None:
            break

    if matched_lines is None:
        return False

    content_runs = [run for run in paragraph.runs if _normalized_for_compare(run.text)]
    groups: list[list] = []
    cursor = 0
    for source_line in matched_lines:
        target_signature = _character_signature(source_line)
        current_signature = Counter()
        group = []
        while cursor < len(content_runs):
            run = content_runs[cursor]
            cursor += 1
            group.append(run)
            current_signature.update(_character_signature(run.text))
            if current_signature == target_signature:
                break
        if current_signature != target_signature:
            return False
        groups.append(group)

    if cursor != len(content_runs):
        return False

    for run in paragraph.runs:
        run.text = ""
    for index, (source_line, group) in enumerate(zip(matched_lines, groups)):
        group[0].text = source_line + ("\n" if index < len(groups) - 1 else "")
    return True


def _replace_paragraph_text(paragraph: Paragraph, text: str) -> None:
    if paragraph.runs:
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(text)


def _restore_missing_totals(document, totals: list[str]) -> None:
    paragraphs = list(_iter_all_paragraphs(document))
    document_text = "\n".join(paragraph.text for paragraph in paragraphs)
    normalized_document = _normalized_for_compare(document_text)

    for total in totals:
        normalized_total = _normalized_for_compare(total)
        if normalized_total and normalized_total in normalized_document:
            continue

        # pdf2docx commonly retains only the trailing currency word when the
        # numeric total overlaps the table boundary.  Reuse that paragraph so
        # the recovered total stays in the correct place instead of moving to
        # the end of the document.
        replacement = None
        for paragraph in paragraphs:
            normalized_paragraph = _normalized_for_compare(paragraph.text)
            if (
                normalized_paragraph
                and len(normalized_paragraph) <= 12
                and normalized_total.endswith(normalized_paragraph)
            ):
                replacement = paragraph
                break

        if replacement is None:
            replacement = document.add_paragraph()
            paragraphs.append(replacement)

        _replace_paragraph_text(replacement, total)
        _mark_paragraph_rtl(replacement)
        normalized_document += normalized_total


def repair_arabic_docx(pdf_path: str, docx_path: str) -> None:
    """Repair one converted DOCX in place using its source PDF."""
    numeric_replacements = extract_numeric_replacements(pdf_path)
    logical_lines = extract_logical_lines(pdf_path)
    totals = [
        line
        for line in logical_lines
        if ARABIC_NUMBER_RE.search(line) and any(word in line for word in TOTAL_WORDS)
    ]
    document = Document(docx_path)

    for paragraph in _iter_all_paragraphs(document):
        for run in paragraph.runs:
            run.text = repair_extracted_text(run.text, numeric_replacements)

    for paragraph in document.paragraphs:
        _align_paragraph_to_source_lines(paragraph, logical_lines)

    for paragraph in _iter_all_paragraphs(document):
        _mark_paragraph_rtl(paragraph)

    for table in document.tables:
        _mark_table_rtl(table)

    _restore_missing_totals(document, totals)

    temporary_path = f"{docx_path}.{uuid.uuid4().hex}.tmp"
    try:
        document.save(temporary_path)
        os.replace(temporary_path, docx_path)
    finally:
        if os.path.exists(temporary_path):
            os.remove(temporary_path)
