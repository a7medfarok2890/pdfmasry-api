"""Google Cloud Document AI integration — optional quality upgrade for the
Arabic PDF→Word and PDF→Excel converters.

This module is fully opt-in and fails soft: every entry point raises
``DocAIUnavailable`` on *any* problem (missing/invalid credentials, network
error, quota exhausted, malformed response, unsupported page count, ...).
Callers (main.py) catch that and fall back to the existing self-hosted /
Adobe pipeline — a user must never see a Document AI-specific error, and
the tool must keep working exactly as before if the env vars below are
never set.

Required environment variables (all five must be set, or ``is_configured()``
returns False and callers skip this module entirely):

  GOOGLE_APPLICATION_CREDENTIALS_JSON  — the full service-account JSON as a
                                          string (NOT a file path). Parsed
                                          straight into memory since
                                          serverless/container filesystems
                                          aren't guaranteed persistent.
  GOOGLE_CLOUD_PROJECT_ID
  GOOGLE_CLOUD_LOCATION                — e.g. "us"
  GOOGLE_DOCAI_OCR_PROCESSOR_ID        — "Document OCR" processor, used by
                                          PDF→Word for text
  GOOGLE_DOCAI_FORM_PROCESSOR_ID       — "Form Parser" processor, used by
                                          PDF→Excel for everything, and by
                                          PDF→Word for table detection
                                          (Document OCR alone frequently
                                          reports no tables on tabular
                                          documents)

Optional:
  GOOGLE_DOCAI_TIMEOUT_SECONDS         — per-request timeout (default 60)
"""
from __future__ import annotations

import json
import os
import threading
from typing import Optional

# Bump this whenever the DOCX/XLSX-building logic below changes format —
# main.py's cache keys on (provider_name, provider_version), so an
# unbumped version means a cached result from before the change gets
# served forever instead of a freshly regenerated one.
DOCAI_PROVIDER_VERSION = "docai-v7-2026-08"
DOCAI_TIMEOUT_SECONDS = int(os.environ.get("GOOGLE_DOCAI_TIMEOUT_SECONDS", "60"))

# Document AI's synchronous processDocument call is capped at 15 pages for
# these processor types. Larger files would need the async batch API (GCS
# input/output buckets), which isn't provisioned here — so we bail out to
# the fallback engine instead of trying and failing mid-request.
MAX_SYNC_PAGES = 15


class DocAIUnavailable(Exception):
    """Any condition that should trigger fallback to the existing engine."""


def _env(name: str) -> Optional[str]:
    val = os.environ.get(name)
    return val if val else None


def is_configured() -> bool:
    """True only if every required env var is present. Cheap — safe to call
    on every request; this is what lets the tools keep working unmodified
    when the credentials haven't been added yet."""
    return all([
        _env("GOOGLE_APPLICATION_CREDENTIALS_JSON"),
        _env("GOOGLE_CLOUD_PROJECT_ID"),
        _env("GOOGLE_CLOUD_LOCATION"),
        _env("GOOGLE_DOCAI_OCR_PROCESSOR_ID"),
        _env("GOOGLE_DOCAI_FORM_PROCESSOR_ID"),
    ])


_client_lock = threading.Lock()
_client_cache: dict[str, object] = {}


def _get_client():
    """Lazily build (and cache for the process lifetime) a
    DocumentProcessorServiceClient from the JSON in
    GOOGLE_APPLICATION_CREDENTIALS_JSON. Never touches disk."""
    with _client_lock:
        cached = _client_cache.get("client")
        if cached is not None:
            return cached

        try:
            from google.cloud import documentai
            from google.oauth2 import service_account
        except ImportError as exc:
            raise DocAIUnavailable(f"مكتبة Document AI غير مثبتة: {exc}")

        raw = _env("GOOGLE_APPLICATION_CREDENTIALS_JSON")
        try:
            info = json.loads(raw)
        except (json.JSONDecodeError, TypeError) as exc:
            raise DocAIUnavailable(f"GOOGLE_APPLICATION_CREDENTIALS_JSON غير صالح: {exc}")

        try:
            credentials = service_account.Credentials.from_service_account_info(
                info, scopes=["https://www.googleapis.com/auth/cloud-platform"],
            )
            location = _env("GOOGLE_CLOUD_LOCATION")
            client = documentai.DocumentProcessorServiceClient(
                credentials=credentials,
                client_options={"api_endpoint": f"{location}-documentai.googleapis.com"},
            )
        except Exception as exc:
            raise DocAIUnavailable(f"فشل تهيئة عميل Document AI: {exc}")

        _client_cache["client"] = client
        return client


def _processor_name(processor_id: str) -> str:
    project = _env("GOOGLE_CLOUD_PROJECT_ID")
    location = _env("GOOGLE_CLOUD_LOCATION")
    return f"projects/{project}/locations/{location}/processors/{processor_id}"


def _guard_page_count(pdf_path: str) -> None:
    try:
        import fitz
        with fitz.open(pdf_path) as doc:
            if doc.page_count > MAX_SYNC_PAGES:
                raise DocAIUnavailable(
                    f"الملف {doc.page_count} صفحة — يتجاوز حد المعالجة الفورية "
                    f"لـ Document AI ({MAX_SYNC_PAGES} صفحة)"
                )
    except DocAIUnavailable:
        raise
    except Exception:
        # Can't even open it to count pages — let the real API call fail
        # naturally below and surface as DocAIUnavailable there.
        pass


def _process_document(pdf_path: str, processor_id: str):
    from google.api_core import exceptions as google_exceptions
    from google.cloud import documentai

    client = _get_client()
    name = _processor_name(processor_id)

    with open(pdf_path, "rb") as f:
        content = f.read()

    request = documentai.ProcessRequest(
        name=name,
        raw_document=documentai.RawDocument(content=content, mime_type="application/pdf"),
    )

    try:
        result = client.process_document(request=request, timeout=DOCAI_TIMEOUT_SECONDS)
    except google_exceptions.ResourceExhausted as exc:
        raise DocAIUnavailable(f"تجاوز الحصة المجانية لـ Document AI: {exc}")
    except google_exceptions.GoogleAPICallError as exc:
        raise DocAIUnavailable(f"خطأ Document AI: {exc}")
    except Exception as exc:
        raise DocAIUnavailable(f"فشل استدعاء Document AI: {exc}")

    return result.document


def _layout_text(document_text: str, layout) -> str:
    """Resolve a Document AI Layout's text_anchor into the substring of the
    document's full text it refers to. Document AI already returns this in
    correct logical reading order (including for Arabic), which is the
    whole reason this module exists — no bidi post-processing needed here,
    unlike the self-hosted pdf2docx/pdfplumber path."""
    if layout is None or not layout.text_anchor.text_segments:
        return ""
    parts = []
    for seg in layout.text_anchor.text_segments:
        start = int(seg.start_index) if seg.start_index else 0
        end = int(seg.end_index)
        parts.append(document_text[start:end])
    return "".join(parts).strip()


def _bbox_from_layout(layout) -> Optional[tuple[float, float, float, float]]:
    """Return (x_min, y_min, x_max, y_max) in normalized [0,1] page
    coordinates for a Layout's bounding_poly, or None if unavailable.

    Normalized coordinates are page-geometry-based, not tied to either
    processor's own text numbering — which is what makes them usable to
    compare layout elements that came from two *different* Document AI
    responses (Document OCR's paragraphs vs. Form Parser's tables) run
    against the same PDF pages.
    """
    if layout is None:
        return None
    poly = layout.bounding_poly
    verts = list(poly.normalized_vertices) if poly and poly.normalized_vertices else []
    if not verts:
        return None
    xs = [v.x for v in verts]
    ys = [v.y for v in verts]
    return (min(xs), min(ys), max(xs), max(ys))


def _bbox_overlaps(a, b, threshold: float = 0.5) -> bool:
    """True if bbox `a` has at least `threshold` fraction of its own area
    covered by bbox `b`. Used to decide whether an OCR paragraph's region
    is already covered by a detected table (so its text is skipped rather
    than duplicated). Missing bounding boxes never count as overlapping —
    better to risk a rare duplicate line than silently drop text."""
    if a is None or b is None:
        return False
    ax0, ay0, ax1, ay1 = a
    bx0, by0, bx1, by1 = b
    iw = max(0.0, min(ax1, bx1) - max(ax0, bx0))
    ih = max(0.0, min(ay1, by1) - max(ay0, by0))
    a_area = max(1e-9, (ax1 - ax0) * (ay1 - ay0))
    return (iw * ih / a_area) >= threshold


def _set_explicit_table_borders(docx_table) -> None:
    """Set visible single-line borders directly in the table's own XML,
    rather than relying solely on the "Table Grid" style name resolving
    correctly in every Word/LibreOffice version that opens the file —
    belt-and-suspenders so the table is never rendered borderless.

    Border weight (sz=6, i.e. 0.75pt) and zeroed cell margins match a
    reference iLovePDF-generated docx byte-for-byte, compared directly
    via its raw XML."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    tbl_pr = docx_table._tbl.tblPr

    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        edge_el = OxmlElement(f"w:{edge}")
        edge_el.set(qn("w:val"), "single")
        edge_el.set(qn("w:sz"), "6")
        edge_el.set(qn("w:space"), "0")
        edge_el.set(qn("w:color"), "000000")
        borders.append(edge_el)
    tbl_pr.append(borders)

    cell_margins = OxmlElement("w:tblCellMar")
    for edge in ("top", "left", "bottom", "right"):
        margin_el = OxmlElement(f"w:{edge}")
        margin_el.set(qn("w:w"), "0")
        margin_el.set(qn("w:type"), "dxa")
        cell_margins.append(margin_el)
    tbl_pr.append(cell_margins)


def _add_docx_table(docx_doc, full_text: str, table, table_style) -> None:
    """Append a real Word table (w:tbl) built from a Document AI table,
    with the same RTL style used everywhere else in the document and
    explicit visible borders."""
    import arabic_docx

    rows = list(table.header_rows) + list(table.body_rows)
    col_count = max((len(row.cells) for row in rows), default=0)
    if not rows or col_count == 0:
        return

    docx_table = docx_doc.add_table(rows=0, cols=col_count)
    docx_table.style = "Table Grid"
    _set_explicit_table_borders(docx_table)
    for row in rows:
        row_cells = docx_table.add_row().cells
        for i, cell in enumerate(row.cells):
            if i >= col_count:
                break
            row_cells[i].text = _layout_text(full_text, cell.layout) if cell.layout else ""
            for paragraph in row_cells[i].paragraphs:
                paragraph.style = table_style
                arabic_docx.mark_paragraph_rtl(
                    paragraph,
                    force=True,
                    align_to_start=True,
                )
    arabic_docx.mark_table_rtl(docx_table, mirror_columns=False)


def _set_section_rtl(docx_doc) -> None:
    """Mark the document's section itself as right-to-left, in addition
    to each paragraph's own style. Every paragraph in this document
    already carries bidi + right alignment via its style, which is the
    correct per-paragraph setting — this adds the document-level flag on
    top, matching how a fully-RTL document is described end to end
    rather than relying solely on per-paragraph settings applying before
    any first render."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    sect_pr = docx_doc.sections[0]._sectPr
    bidi_el = sect_pr.find(qn("w:bidi"))
    if bidi_el is None:
        sect_pr.append(OxmlElement("w:bidi"))


def _build_docx_from_ocr(ocr_document, form_document, output_path: str) -> None:
    """Build a DOCX from Document OCR's text merged with Form Parser's
    tables, preserving both running text and real table structure.

    Document OCR gives excellent Arabic text quality but frequently
    reports no page.tables at all for documents that are visually
    tabular (e.g. financial reports) — it's a text processor first.
    Form Parser is the processor actually built for table detection, so
    its tables are used here instead, keyed onto the OCR text by page
    position (bounding boxes) rather than text offset, since the two
    responses are independent documents with unrelated text numbering.

    ``form_document`` may be None (its call failed or was skipped) — the
    conversion still proceeds as OCR-text-only in that case rather than
    losing the whole result over a table-detection hiccup.
    """
    from docx import Document as DocxDocument

    import arabic_docx

    docx_doc = DocxDocument()
    _set_section_rtl(docx_doc)

    # RTL alignment is baked into these styles (jc="left" + bidi in the
    # style definition itself) and repeated as direct paragraph formatting,
    # not decided per-paragraph from that
    # paragraph's own text — this document is RTL by construction end to
    # end, so a cell whose whole content is "100%" or "6,993,250" must
    # still render right-aligned, matching how other Word-generation
    # tools (e.g. iLovePDF) structure this rather than skipping alignment
    # on paragraphs with no Arabic characters of their own.
    body_style = arabic_docx.ensure_rtl_paragraph_style(docx_doc, "DocAI Body Text", font_name="Arial")
    table_style = arabic_docx.ensure_rtl_paragraph_style(
        docx_doc, "DocAI Table Paragraph", font_name="Arial", spacing_before_twips=14,
    )

    ocr_text = ocr_document.text or ""
    form_text = form_document.text if form_document is not None else ""
    wrote_any = False
    ocr_pages = list(ocr_document.pages)
    form_pages = list(form_document.pages) if form_document is not None else []

    for page_index, page in enumerate(ocr_pages):
        paragraphs = list(page.paragraphs or [])
        form_page = form_pages[page_index] if page_index < len(form_pages) else None
        tables = list(form_page.tables or []) if form_page is not None else []
        table_bboxes = [_bbox_from_layout(t.layout) for t in tables]

        if paragraphs or tables:
            blocks: list[tuple[float, str, object]] = [
                (bbox[1] if bbox else 0.0, "table", t) for t, bbox in zip(tables, table_bboxes)
            ]
            for para in paragraphs:
                p_bbox = _bbox_from_layout(para.layout)
                if any(_bbox_overlaps(p_bbox, tb) for tb in table_bboxes):
                    continue  # region already covered by a detected table — skip to avoid duplicating it as text
                text = _layout_text(ocr_text, para.layout)
                if text.strip():
                    blocks.append((p_bbox[1] if p_bbox else 0.0, "paragraph", text))
            blocks.sort(key=lambda b: b[0])

            for _sort_key, kind, obj in blocks:
                if kind == "table":
                    _add_docx_table(docx_doc, form_text, obj, table_style)
                else:
                    paragraph = docx_doc.add_paragraph(obj, style=body_style)
                    arabic_docx.mark_paragraph_rtl(
                        paragraph,
                        force=True,
                        align_to_start=True,
                    )
                wrote_any = True
        elif page.layout:
            for line in _layout_text(ocr_text, page.layout).splitlines():
                if not line.strip():
                    continue
                paragraph = docx_doc.add_paragraph(line, style=body_style)
                arabic_docx.mark_paragraph_rtl(
                    paragraph,
                    force=True,
                    align_to_start=True,
                )
                wrote_any = True

        if page_index < len(ocr_pages) - 1:
            docx_doc.add_page_break()

    if not wrote_any:
        raise DocAIUnavailable("Document AI لم يستخرج أي نص من الملف")

    docx_doc.save(output_path)


def _build_xlsx_from_form(document, output_path: str) -> None:
    from openpyxl import Workbook

    wb = Workbook()
    wb.remove(wb.active)
    full_text = document.text or ""
    wrote_any = False

    for page_num, page in enumerate(document.pages, start=1):
        ws = wb.create_sheet(title=f"Page {page_num}"[:31])
        tables = list(page.tables or [])
        if tables:
            for table in tables:
                for row in list(table.header_rows) + list(table.body_rows):
                    cells = [_layout_text(full_text, cell.layout) for cell in row.cells]
                    if any(cells):
                        ws.append(cells)
                        wrote_any = True
                ws.append([])  # blank row between tables, matches self-hosted layout
        else:
            page_text = _layout_text(full_text, page.layout) if page.layout else ""
            for line in page_text.splitlines():
                if line.strip():
                    ws.append([line])
                    wrote_any = True

    if not wb.sheetnames:
        wb.create_sheet(title="Empty")
    if not wrote_any:
        raise DocAIUnavailable("Document AI لم يستخرج أي جداول أو نص من الملف")

    wb.save(output_path)


def process_pdf_to_docx(input_path: str, output_path: str) -> None:
    """PDF → DOCX via Document AI.

    Calls the Document OCR processor for text (the source of the confirmed
    Arabic quality improvement) and additionally calls Form Parser for
    table detection — Document OCR alone often reports no tables at all
    for visually tabular documents like financial reports, since it isn't
    the processor built for that. Only a Document OCR failure raises
    DocAIUnavailable; if the extra Form Parser call fails, the conversion
    still proceeds as text-only rather than losing an otherwise-good
    result over a table-detection hiccup.

    Raises DocAIUnavailable on a Document OCR failure — the caller must
    catch this and fall back to the self-hosted/Adobe pipeline rather than
    propagate it.
    """
    try:
        _guard_page_count(input_path)
        ocr_document = _process_document(input_path, _env("GOOGLE_DOCAI_OCR_PROCESSOR_ID"))
        try:
            form_document = _process_document(input_path, _env("GOOGLE_DOCAI_FORM_PROCESSOR_ID"))
        except DocAIUnavailable:
            form_document = None
        _build_docx_from_ocr(ocr_document, form_document, output_path)
    except DocAIUnavailable:
        raise
    except Exception as exc:
        raise DocAIUnavailable(f"خطأ غير متوقع أثناء معالجة Document AI: {exc}")

    if not os.path.exists(output_path) or os.path.getsize(output_path) == 0:
        raise DocAIUnavailable("لم يُنشئ Document AI ملف Word صالح")


def process_pdf_to_xlsx(input_path: str, output_path: str) -> None:
    """PDF → XLSX via Document AI's Form Parser processor.

    Raises DocAIUnavailable on any failure — the caller must catch this and
    fall back to the self-hosted/Adobe pipeline rather than propagate it.
    """
    try:
        _guard_page_count(input_path)
        document = _process_document(input_path, _env("GOOGLE_DOCAI_FORM_PROCESSOR_ID"))
        _build_xlsx_from_form(document, output_path)
    except DocAIUnavailable:
        raise
    except Exception as exc:
        raise DocAIUnavailable(f"خطأ غير متوقع أثناء معالجة Document AI: {exc}")

    if not os.path.exists(output_path) or os.path.getsize(output_path) == 0:
        raise DocAIUnavailable("لم يُنشئ Document AI ملف Excel صالح")


def _reset_client_cache_for_testing() -> None:
    """Used by unit tests to force _get_client() to rebuild. NOT wired to
    any production path."""
    with _client_lock:
        _client_cache.pop("client", None)
