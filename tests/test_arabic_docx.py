"""Regression tests for Arabic PDF-to-Word post-processing."""
from __future__ import annotations

from docx import Document
from docx.oxml.ns import qn

import arabic_docx


def test_repairs_browser_pdf_combining_mark_artifacts():
    damaged = "اختبار الحروف العربية والتشكيل: ُمؤَّسَّسسة الصيانة السريعة."

    assert arabic_docx.repair_arabic_combining_marks(damaged) == (
        "اختبار الحروف العربية والتشكيل: مُؤسَّسة الصيانة السريعة."
    )


def test_repairs_only_numbers_confirmed_by_source_coordinates():
    replacements = {
        "١٢/٨٠/٦٢٠٢": "٢٠٢٦/٠٨/٢١",
        "٠٥٫٥٢١": "١٢٥٫٥٠",
        "٠٠٫٦٢١٬١": "١٬١٢٦٫٠٠",
    }

    assert arabic_docx.repair_extracted_text(
        "التاريخ: ١٢/٨٠/٦٢٠٢ — السعر ٠٥٫٥٢١",
        replacements,
    ) == "التاريخ: ٢٠٢٦/٠٨/٢١ — السعر ١٢٥٫٥٠"

    # A token that is not proven reversed by source coordinates is untouched.
    assert arabic_docx.repair_extracted_text("رقم ٢٠٢٦", replacements) == "رقم ٢٠٢٦"


def test_repair_docx_adds_rtl_and_restores_missing_total(tmp_path, monkeypatch):
    docx_path = tmp_path / "converted.docx"
    source_path = tmp_path / "source.pdf"
    source_path.write_bytes(b"%PDF-test-fixture")

    document = Document()
    document.add_paragraph("العميل: أحمد — التاريخ: ١٢/٨٠/٦٢٠٢")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "الصنف"
    table.cell(0, 1).text = "الإجمالي"
    table.cell(1, 0).text = "فلتر زيت"
    table.cell(1, 1).text = "٠٠٫١٥٢ / ٠٥٫٥٢١"
    document.add_paragraph(" جنيه.")
    document.save(docx_path)

    monkeypatch.setattr(
        arabic_docx,
        "extract_numeric_replacements",
        lambda _path: {
            "١٢/٨٠/٦٢٠٢": "٢٠٢٦/٠٨/٢١",
            "٠٠٫١٥٢": "٢٥١٫٠٠",
            "٠٥٫٥٢١": "١٢٥٫٥٠",
        },
    )
    monkeypatch.setattr(
        arabic_docx,
        "extract_logical_lines",
        lambda _path: ["الإجمالي المستحق: ١٬١٢٦٫٠٠ جنيه."],
    )

    arabic_docx.repair_arabic_docx(str(source_path), str(docx_path))

    repaired = Document(docx_path)
    text = "\n".join(
        [paragraph.text for paragraph in repaired.paragraphs]
        + [cell.text for row in repaired.tables[0].rows for cell in row.cells]
    )
    assert "٢٠٢٦/٠٨/٢١" in text
    assert "٢٥١٫٠٠" in text
    assert "١٢٥٫٥٠" in text
    assert "الإجمالي المستحق: ١٬١٢٦٫٠٠ جنيه." in text
    assert "١٢/٨٠/٦٢٠٢" not in text
    assert "٠٠٫١٥٢" not in text
    assert "٠٥٫٥٢١" not in text

    arabic_paragraph = repaired.paragraphs[0]
    assert arabic_paragraph._p.pPr.find(qn("w:bidi")) is not None
    assert arabic_paragraph.runs[0]._r.rPr.find(qn("w:rtl")) is not None
    assert repaired.tables[0]._tbl.tblPr.find(qn("w:bidiVisual")) is not None


def test_aligns_concatenated_and_physically_ordered_source_lines():
    document = Document()
    paragraph = document.add_paragraph()
    paragraph.add_run("فاتورة تجريبية PDF مصري")
    paragraph.add_run("فاتورة تجريبية")
    paragraph.add_run("\n")
    paragraph.add_run("٢٠٢٦/٠٨/٢١")
    paragraph.add_run(" :العميل: أحمد محمد — التاريخ")

    matched = arabic_docx._align_paragraph_to_source_lines(
        paragraph,
        [
            "فاتورة تجريبية PDF مصري",
            "فاتورة تجريبية",
            "العميل: أحمد محمد — التاريخ: ٢٠٢٦/٠٨/٢١",
        ],
    )

    assert matched is True
    assert paragraph.text == (
        "فاتورة تجريبية PDF مصري\n"
        "فاتورة تجريبية\n"
        "العميل: أحمد محمد — التاريخ: ٢٠٢٦/٠٨/٢١"
    )


def test_force_rtl_matches_ilovepdf_direct_paragraph_properties():
    document = Document()
    paragraph = document.add_paragraph()
    arabic_run = paragraph.add_run("التاريخ: ")
    latin_run = paragraph.add_run("14 February 2026")

    arabic_docx.mark_paragraph_rtl(
        paragraph,
        force=True,
        align_to_start=True,
    )

    p_pr = paragraph._p.pPr
    assert p_pr.find(qn("w:bidi")).get(qn("w:val")) == "1"
    assert p_pr.find(qn("w:jc")).get(qn("w:val")) == "left"
    assert arabic_run._r.rPr.find(qn("w:rtl")).get(qn("w:val")) == "1"
    assert latin_run._r.rPr is None or latin_run._r.rPr.find(qn("w:rtl")) is None


def test_docai_table_keeps_logical_column_order():
    document = Document()
    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "التحليل"
    table.cell(0, 1).text = "البيان المالي"

    # Simulate a stale/misapplied table flag, then ensure the Document AI
    # path removes it instead of mirroring already-logical columns twice.
    arabic_docx.mark_table_rtl(table)
    assert table._tbl.tblPr.find(qn("w:bidiVisual")) is not None

    arabic_docx.mark_table_rtl(table, mirror_columns=False)
    assert table._tbl.tblPr.find(qn("w:bidiVisual")) is None
    assert [cell.text for cell in table.rows[0].cells] == [
        "التحليل",
        "البيان المالي",
    ]
